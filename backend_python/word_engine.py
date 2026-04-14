"""
Word 教案生成引擎
=================
将 LLM 输出的结构化 JSON 渲染为专业排版的 Word 教案文档 (.docx)。

教案固定包含以下五大模块：
  1. 教学目标
  2. 教学过程
  3. 教学方法
  4. 课堂活动设计
  5. 课后作业
"""

import os
import json
import uuid
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(CURRENT_DIR)
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "ppts")  # 复用已有的下载目录

# 模块中文标题映射（固定五大模块）
SECTION_KEYS = [
    ("teaching_objectives",  "一、教学目标"),
    ("teaching_process",     "二、教学过程"),
    ("teaching_methods",     "三、教学方法"),
    ("activity_design",      "四、课堂活动设计"),
    ("homework",             "五、课后作业"),
]


# ==========================================
# LLM 提示词：将用户主题转化为结构化教案 JSON
# ==========================================

LESSON_PLAN_SYSTEM_PROMPT = """你是一位拥有丰富教学经验的资深教育工作者。你的任务是根据用户提供的教学主题，生成一份详细的教案文档内容。

=== 输出格式要求（必须严格遵守）===
输出必须是一个 JSON 对象，包含以下字段：
{
  "title": "教案标题（如：《望庐山瀑布》教学教案）",
  "grade_info": "适用年级/学段（如：小学四年级语文）",
  "teaching_objectives": [
    "目标1：...",
    "目标2：...",
    "目标3：..."
  ],
  "teaching_process": [
    {
      "stage": "阶段名称（如：导入新课）",
      "duration": "时长（如：5分钟）",
      "content": "详细描述该阶段的教学活动内容"
    },
    {
      "stage": "阶段名称",
      "duration": "时长",
      "content": "详细描述"
    }
  ],
  "teaching_methods": [
    "方法1：...",
    "方法2：..."
  ],
  "activity_design": [
    {
      "name": "活动名称",
      "description": "活动详细说明"
    }
  ],
  "homework": [
    "作业1：...",
    "作业2：..."
  ]
}

=== 内容质量要求 ===
1. 教学目标：要涵盖知识与技能、过程与方法、情感态度与价值观三个维度。
2. 教学过程：至少包含导入、新授、练习、小结四个阶段，每个阶段要有具体活动描述。
3. 教学方法：结合主题选用合适的教学方法（讲授法、讨论法、探究法、情境教学法等）。
4. 课堂活动设计：至少2个有创意的课堂互动活动。
5. 课后作业：2-3项，兼顾巩固与拓展。

只输出纯 JSON，不要加 markdown 标记或解释文字。"""


def build_lesson_plan_prompt(topic):
    """构建教案生成的 LLM 请求消息"""
    return [
        {"role": "system", "content": LESSON_PLAN_SYSTEM_PROMPT},
        {"role": "user", "content": f"请为以下教学主题生成一份完整的教案：{topic}"}
    ]


def _set_cell_shading(cell, color_hex):
    """为表格单元格设置底色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading_elm.append(shading)


def _add_styled_heading(doc, text, level=2):
    """添加带有自定义样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)  # 深红色标题
        run.font.bold = True
    return heading


def _add_body_paragraph(doc, text, indent=True, bold=False):
    """添加正文段落（宋体风格，1.5倍行距）"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold

    para_format = para.paragraph_format
    para_format.space_after = Pt(6)
    para_format.line_spacing = 1.5

    if indent:
        para_format.first_line_indent = Cm(0.74)  # 两个字符缩进

    return para


def _add_list_item(doc, text, bullet_char="•"):
    """添加带项目符号的列表项"""
    para = doc.add_paragraph()
    run = para.add_run(f"{bullet_char} {text}")
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    para_format = para.paragraph_format
    para_format.space_after = Pt(4)
    para_format.line_spacing = 1.5
    para_format.left_indent = Cm(1)

    return para


# ==========================================
# 核心渲染函数
# ==========================================

def generate_lesson_plan_docx(lesson_data, output_path=None):
    """
    将结构化教案 JSON 渲染成 Word 文档。

    参数:
        lesson_data: dict, LLM 生成的教案 JSON 数据
        output_path: str, 可选，输出文件路径。留空则自动生成。

    返回:
        str: 生成的 .docx 文件完整路径
    """
    if output_path is None:
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        filename = f"LessonPlan_{uuid.uuid4().hex[:8]}.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)

    doc = Document()

    # ===== 全局字体默认值 =====
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 1. 文档标题 =====
    title_text = lesson_data.get("title", "教学教案")
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title_text)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # ===== 2. 基本信息行 =====
    grade_info = lesson_data.get("grade_info", "")
    if grade_info:
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(f"适用范围：{grade_info}")
        info_run.font.size = Pt(11)
        info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 分隔线
    doc.add_paragraph("─" * 50)

    # ===== 3. 教学目标 =====
    _add_styled_heading(doc, "一、教学目标")
    objectives = lesson_data.get("teaching_objectives", [])
    if isinstance(objectives, list):
        for i, obj in enumerate(objectives, 1):
            _add_list_item(doc, obj, bullet_char=f"{i}.")
    else:
        _add_body_paragraph(doc, str(objectives))

    # ===== 4. 教学过程 =====
    _add_styled_heading(doc, "二、教学过程")
    process = lesson_data.get("teaching_process", [])
    if isinstance(process, list):
        for step in process:
            if isinstance(step, dict):
                stage = step.get("stage", "")
                duration = step.get("duration", "")
                content = step.get("content", "")

                # 阶段标题（加粗）
                stage_header = f"【{stage}】"
                if duration:
                    stage_header += f"（{duration}）"
                _add_body_paragraph(doc, stage_header, indent=False, bold=True)

                # 阶段内容
                if content:
                    _add_body_paragraph(doc, content)
            else:
                _add_list_item(doc, str(step))
    else:
        _add_body_paragraph(doc, str(process))

    # ===== 5. 教学方法 =====
    _add_styled_heading(doc, "三、教学方法")
    methods = lesson_data.get("teaching_methods", [])
    if isinstance(methods, list):
        for method in methods:
            _add_list_item(doc, method)
    else:
        _add_body_paragraph(doc, str(methods))

    # ===== 6. 课堂活动设计 =====
    _add_styled_heading(doc, "四、课堂活动设计")
    activities = lesson_data.get("activity_design", [])
    if isinstance(activities, list):
        for act in activities:
            if isinstance(act, dict):
                name = act.get("name", "")
                desc = act.get("description", "")
                _add_body_paragraph(doc, f"📌 {name}", indent=False, bold=True)
                if desc:
                    _add_body_paragraph(doc, desc)
            else:
                _add_list_item(doc, str(act))
    else:
        _add_body_paragraph(doc, str(activities))

    # ===== 7. 课后作业 =====
    _add_styled_heading(doc, "五、课后作业")
    homework = lesson_data.get("homework", [])
    if isinstance(homework, list):
        for i, hw in enumerate(homework, 1):
            _add_list_item(doc, hw, bullet_char=f"{i}.")
    else:
        _add_body_paragraph(doc, str(homework))

    # ===== 保存文件 =====
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"[OK] Word 教案文件已保存: {output_path}")

    return output_path
